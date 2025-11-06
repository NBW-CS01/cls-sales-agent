"""
NDA Generator Lambda Function
Generates NDAs by fetching company data from Companies House and populating templates
"""

import json
import boto3
import os
from typing import Dict, List
from datetime import datetime
import io
import re
from companies_house import CompaniesHouseClient

# Try to import python-docx, will be available in Lambda layer
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("Warning: python-docx not available. Will fail when generating documents.")

s3 = boto3.client('s3')

KNOWLEDGE_BASE_BUCKET = os.environ.get('KNOWLEDGE_BASE_BUCKET', '')
COMPANIES_HOUSE_API_KEY = os.environ.get('COMPANIES_HOUSE_API_KEY', '')  # Optional
NDA_TEMPLATE_KEY = 'templates/Non-Disclosure Agreement Template v01 JC OFFLINE.docx'


def lambda_handler(event, context):
    """
    Main handler for Bedrock Agent action group - NDA generation
    """
    print(f"Received event: {json.dumps(event)}")

    action_group = event.get('actionGroup')
    function = event.get('function')
    parameters = event.get('parameters', [])

    # Extract parameters
    company_identifier = None
    signatory_name = None
    signatory_title = None

    for param in parameters:
        if param['name'] == 'company_identifier':
            company_identifier = param['value']
        elif param['name'] == 'signatory_name':
            signatory_name = param['value']
        elif param['name'] == 'signatory_title':
            signatory_title = param['value']

    if not company_identifier:
        return error_response("Company identifier (name or number) is required")

    if not signatory_name or not signatory_title:
        return error_response("Signatory name and title are required")

    try:
        # Generate the NDA
        result = generate_nda(
            company_identifier=company_identifier,
            signatory_name=signatory_name,
            signatory_title=signatory_title
        )

        # Format response for Bedrock Agent
        response_body = {
            "TEXT": {
                "body": json.dumps(result)
            }
        }

        action_response = {
            'actionGroup': action_group,
            'function': function,
            'functionResponse': {
                'responseBody': response_body
            }
        }

        return {
            'response': action_response,
            'messageVersion': event['messageVersion']
        }

    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return error_response(str(e))


def generate_nda(company_identifier: str, signatory_name: str, signatory_title: str) -> Dict:
    """
    Generate an NDA document

    Args:
        company_identifier: Company name or Companies House number
        signatory_name: Name of the person signing
        signatory_title: Title/position of the signatory

    Returns:
        Dictionary with generation results
    """
    print(f"Generating NDA for: {company_identifier}")

    # Step 1: Get company details from Companies House
    ch_client = CompaniesHouseClient(COMPANIES_HOUSE_API_KEY or None)
    company_data = None

    # Check if it's a company number or name
    if is_company_number(company_identifier):
        print(f"Fetching company by number: {company_identifier}")
        company_data = ch_client.get_company_details(company_identifier)
    else:
        print(f"Searching for company by name: {company_identifier}")
        # Search by name and use first result
        from companies_house import search_company_by_name
        search_results = search_company_by_name(company_identifier, COMPANIES_HOUSE_API_KEY or None)

        if not search_results:
            raise ValueError(f"No company found with name: {company_identifier}")

        # Use the first result
        first_result = search_results[0]
        print(f"Found company: {first_result['company_name']} ({first_result['company_number']})")
        company_data = ch_client.get_company_details(first_result['company_number'])

    print(f"Company data retrieved: {company_data['company_name']}")

    # Step 2: Download NDA template from S3
    print(f"Downloading template from S3: {NDA_TEMPLATE_KEY}")
    template_obj = s3.get_object(Bucket=KNOWLEDGE_BASE_BUCKET, Key=NDA_TEMPLATE_KEY)
    template_content = template_obj['Body'].read()

    # Step 3: Populate template with company data
    print("Populating template...")
    populated_doc = populate_nda_template(
        template_content=template_content,
        company_data=company_data,
        signatory_name=signatory_name,
        signatory_title=signatory_title
    )

    # Step 4: Save generated NDA to S3
    timestamp = datetime.utcnow().strftime('%Y%m%d-%H%M%S')
    company_name_safe = re.sub(r'[^a-zA-Z0-9-]', '_', company_data['company_name'])
    output_key = f"generated-ndas/NDA_{company_name_safe}_{timestamp}.docx"

    print(f"Saving NDA to S3: {output_key}")
    s3.put_object(
        Bucket=KNOWLEDGE_BASE_BUCKET,
        Key=output_key,
        Body=populated_doc,
        ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        Metadata={
            'company_name': company_data['company_name'],
            'company_number': company_data['company_number'],
            'signatory_name': signatory_name,
            'signatory_title': signatory_title,
            'generated_date': timestamp
        }
    )

    # Generate download URL (presigned, valid for 1 hour)
    download_url = s3.generate_presigned_url(
        'get_object',
        Params={'Bucket': KNOWLEDGE_BASE_BUCKET, 'Key': output_key},
        ExpiresIn=3600
    )

    return {
        'success': True,
        'message': f'NDA generated successfully for {company_data["company_name"]}',
        'company_details': {
            'name': company_data['company_name'],
            'number': company_data['company_number'],
            'type': company_data['company_type'],
            'jurisdiction': company_data['jurisdiction'],
            'address': company_data['registered_office_address']
        },
        'signatory': {
            'name': signatory_name,
            'title': signatory_title
        },
        's3_key': output_key,
        'download_url': download_url,
        'expires_in': '1 hour'
    }


def populate_nda_template(
    template_content: bytes,
    company_data: Dict,
    signatory_name: str,
    signatory_title: str
) -> bytes:
    """
    Populate NDA template with company and signatory information

    Args:
        template_content: Word document template as bytes
        company_data: Dictionary with company information from Companies House
        signatory_name: Name of signatory
        signatory_title: Title of signatory

    Returns:
        Populated document as bytes
    """
    # Load document from bytes
    doc = Document(io.BytesIO(template_content))

    # Build replacement dictionary
    replacements = {
        '[Company Name]': company_data['company_name'],
        '[Company Number]': company_data['company_number'],
        '[Company Type]': company_data['company_type'],
        '[Registered Office Address]': company_data['registered_office_address'],
        '[Jurisdiction]': company_data['jurisdiction'],
        '[Signatory Name]': signatory_name,
        '[Signatory Title]': signatory_title,
        '[Date]': datetime.utcnow().strftime('%d %B %Y'),
    }

    print(f"Replacements: {replacements}")

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                print(f"Replacing {key} with {value} in paragraph")
                # Replace in runs to preserve formatting
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        print(f"Replacing {key} with {value} in table cell")
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return output.read()


def is_company_number(identifier: str) -> bool:
    """
    Check if the identifier looks like a Companies House number

    Company numbers are typically:
    - 8 characters (with leading zeros)
    - May have letter prefix (e.g., SC, OC, NI)
    - Format: [PREFIX]NNNNNN or NNNNNNNN
    """
    # Remove spaces
    identifier = identifier.replace(' ', '').upper()

    # Check if it matches company number pattern
    # Alphanumeric, 6-8 characters
    if re.match(r'^[A-Z]{0,2}\d{6,8}$', identifier):
        return True

    return False


def error_response(error_message: str) -> Dict:
    """Format error response for Bedrock Agent"""
    return {
        'response': {
            'functionResponse': {
                'responseBody': {
                    'TEXT': {
                        'body': json.dumps({
                            'success': False,
                            'error': error_message
                        })
                    }
                }
            }
        }
    }


# Testing
if __name__ == '__main__':
    # Test the NDA generation locally
    test_event = {
        'agent': 'test',
        'actionGroup': 'NDAs',
        'function': 'generateNDA',
        'parameters': [
            {'name': 'company_identifier', 'value': '01234567'},
            {'name': 'signatory_name', 'value': 'Patrick Godden'},
            {'name': 'signatory_title', 'value': 'Director'}
        ],
        'messageVersion': '1.0'
    }

    result = lambda_handler(test_event, None)
    print(json.dumps(result, indent=2))
